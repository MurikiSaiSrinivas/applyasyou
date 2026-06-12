"""Reusable building blocks for tailored resumes.

Use from a per-resume script:

    from resume_lib import *
    doc = create_document()
    centered(doc, "YOUR NAME", size=20, bold=True)
    ...
    save_docx_and_pdf(doc, OUT_DIR, "YourName_Frontend")

Each helper takes the Document as its first arg so the script can compose
the resume top to bottom in plain Python.

save_docx_and_pdf writes a <basename>.content.md sidecar next to the docx/pdf
capturing the structured content, so the agent can reference what a resume says
without re-rendering the PDF.

Requires: python-docx, docx2pdf, and Microsoft Word (docx2pdf renders the PDF
via Word). On a machine without Word, the .docx is still written; the .pdf step
prints a note and is skipped.
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===== Document setup =====

def create_document(font_name="Calibri", font_size=10,
                    top=0.5, bottom=0.5, left=0.75, right=0.75):
    """Return a Document configured with tight margins and a default font."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    doc._clog = []
    doc._clog_suppress = False
    return doc


# ===== Content recording (sidecar capture) =====

def _rec(doc, kind, **data):
    try:
        if getattr(doc, "_clog_suppress", False):
            return
        log = getattr(doc, "_clog", None)
        if log is None:
            return
        log.append({"kind": kind, **data})
    except Exception:
        pass


class _SuppressClog:
    def __init__(self, doc):
        self.doc = doc
        self.prev = False
    def __enter__(self):
        self.prev = getattr(self.doc, "_clog_suppress", False)
        self.doc._clog_suppress = True
        return self
    def __exit__(self, *a):
        self.doc._clog_suppress = self.prev


# ===== Low-level utilities =====

def add_hr_below(p):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '000000')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ===== Paragraph helpers =====

def centered(doc, text, size=10, bold=False, italic=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def section_heading(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(size)
    add_hr_below(p)
    _rec(doc, "section", text=text)
    return p


def normal_para(doc, text, bold=False, italic=False, size=10,
                space_after=2, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    _rec(doc, "para", text=text, italic=italic, bold=bold)
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for r in list(p.runs):
        r.text = ""
    r = p.add_run(text)
    r.font.size = Pt(size)
    _rec(doc, "bullet", text=text)
    return p


def labeled_bullet(doc, label, body, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for r in list(p.runs):
        r.text = ""
    lr = p.add_run(label)
    lr.bold = True
    lr.font.size = Pt(size)
    br = p.add_run(body)
    br.font.size = Pt(size)
    _rec(doc, "skill", label=label, body=body)
    return p


def two_col_line(doc, left_text, right_text, left_bold=True, size=10,
                 space_before=6, right_tab_in=7.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(right_tab_in), WD_TAB_ALIGNMENT.RIGHT)
    lr = p.add_run(left_text)
    lr.bold = left_bold
    lr.font.size = Pt(size)
    p.add_run("\t")
    rr = p.add_run(right_text)
    rr.font.size = Pt(size)
    return p


# ===== Higher-level blocks =====

def header_block(doc, name, title, contact, name_size=20, title_size=11, contact_size=10):
    _rec(doc, "header", name=name, title=title, contact=contact, links=[])
    with _SuppressClog(doc):
        centered(doc, name, size=name_size, bold=True)
        centered(doc, title, size=title_size)
        centered(doc, contact, size=contact_size, space_after=4)


def job_block(doc, title, dates, company_location, bullets):
    _rec(doc, "job", title=title, dates=dates,
         company_location=company_location, bullets=list(bullets))
    with _SuppressClog(doc):
        two_col_line(doc, title, dates, space_before=6)
        normal_para(doc, company_location, italic=True, space_after=2)
        for b in bullets:
            bullet(doc, b)


def project_block(doc, title, stack_year, bullets):
    _rec(doc, "project", title=title, stack_year=stack_year, bullets=list(bullets))
    with _SuppressClog(doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(title + " ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run("— " + stack_year)
        r2.italic = True
        r2.font.size = Pt(10)
        for b in bullets:
            bullet(doc, b)


def education_block(doc, degree, dates, school_location):
    _rec(doc, "education", degree=degree, dates=dates, school_location=school_location)
    with _SuppressClog(doc):
        two_col_line(doc, degree, dates, space_before=2)
        normal_para(doc, school_location, italic=True, space_after=2)


# ===== Output =====

def save_docx_and_pdf(doc, out_dir, basename, make_pdf=True, make_content=True):
    os.makedirs(out_dir, exist_ok=True)
    docx_path = os.path.join(out_dir, basename + ".docx")
    doc.save(docx_path)
    print(f"DOCX: {docx_path}  ({round(os.path.getsize(docx_path)/1024, 1)} KB)")
    if make_content:
        try:
            content_path = os.path.join(out_dir, basename + ".content.md")
            _write_sidecar(getattr(doc, "_clog", []), content_path, basename, out_dir)
            print(f"CONTENT: {content_path}  ({round(os.path.getsize(content_path)/1024, 1)} KB)")
        except Exception as e:
            print(f"CONTENT: failed to write sidecar ({e})")
    if make_pdf:
        try:
            from docx2pdf import convert
            pdf_path = os.path.join(out_dir, basename + ".pdf")
            convert(docx_path, pdf_path)
            print(f"PDF:  {pdf_path}  ({round(os.path.getsize(pdf_path)/1024, 1)} KB)")
        except ImportError:
            print("PDF: docx2pdf not installed. .docx is ready; open it in Word/Docs to export PDF.")
        except Exception as e:
            print(f"PDF: render failed ({e}). .docx is ready; open it in Word/Docs to export PDF.")
    return docx_path


def _write_sidecar(clog, path, basename, out_dir):
    from datetime import date
    folder = os.path.basename(os.path.normpath(out_dir))
    lines = []
    lines.append(f"## {basename}")
    lines.append(f"_Folder: `{folder}/` · captured {date.today().isoformat()}_")
    lines.append("")
    for ev in clog:
        if ev["kind"] == "header":
            name = ev.get("name", "")
            title = ev.get("title", "")
            contact = ev.get("contact", "")
            lines.append(f"**{name}**  ")
            lines.append(f"_{title}_  ")
            if contact:
                lines.append(contact)
            lines.append("")
            break
    for ev in clog:
        k = ev["kind"]
        if k == "header":
            continue
        if k == "section":
            lines.append("")
            lines.append(f"### {ev['text']}")
            continue
        if k == "para":
            lines.append("")
            lines.append(ev["text"])
            continue
        if k == "skill":
            lines.append(f"- **{ev['label'].rstrip()}** {ev['body']}")
            continue
        if k == "bullet":
            lines.append(f"- {ev['text']}")
            continue
        if k == "job":
            lines.append("")
            lines.append(f"**{ev['title']}** — _{ev['company_location']}_ — {ev['dates']}")
            for b in ev["bullets"]:
                lines.append(f"  - {b}")
            continue
        if k == "project":
            lines.append("")
            lines.append(f"**{ev['title']}** — _{ev['stack_year']}_")
            for b in ev["bullets"]:
                lines.append(f"  - {b}")
            continue
        if k == "education":
            lines.append("")
            lines.append(f"**{ev['degree']}** — _{ev['school_location']}_ — {ev['dates']}")
            continue
    content = "\n".join(lines).rstrip() + "\n"
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


# ===== Hyperlinks =====

def add_hyperlink(paragraph, url, text, size=10, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink
